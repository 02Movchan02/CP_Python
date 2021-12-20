from tkinter import*
from tkinter import ttk
from tkinter import messagebox as mb
import sqlite3 as sql
import docx
from docx.shared import Pt
from datetime import datetime, timedelta

conn=sql.connect('Hotel.db')
cur=conn.cursor()

rootWork = Tk()
rootWork.title("Аренда")

l1 = Label(rootWork, text="Номер комнаты")
cb1 = ttk.Combobox(rootWork, state='readonly')

l2 = Label(rootWork, text="Клиент")
cb2 = ttk.Combobox(rootWork, state='readonly')

l3 = Label(rootWork, text="Количество дней")
e1 = Entry(rootWork)

ds = ""
de = ""

cd = ""

def addCountDay(countDay):
    global cd
    e1.insert(0, countDay)
    cd = countDay

def isChecked():
    global dateStart, dateEnd, lEnd, countDay
    def dateAccept():
        global ds, de
        ds = datetime.strptime(dateStart.get(), "%d-%m-%Y")
        de = datetime.strptime(dateEnd.get(), "%d-%m-%Y")
            
        delta = de-ds
        
        if (delta.days > 0):
            countDay = delta.days
            addCountDay(countDay)
            new_root.destroy()
        else:
            mb.showinfo("Ошибка", "Дата конца не может быть меньше даты начала")
        
    if p1.get()=="Бронь есть":
        # появление поля для ввода даты + подсчёт количества дней
        new_root = Tk()
        titLab = Label(new_root, text = "Выбор даты").pack(side='top', padx=10, pady=5)
        dateStart = Entry(new_root)
        dateStart.pack(side='left', padx=5)
        lEnd = Label(new_root, text="-").pack(side='left', padx=10, pady=5)
        dateEnd = Entry(new_root)
        dateEnd.pack(side='left', padx=5)
        baccDate = Button(new_root, text = "Выбрать", command=dateAccept)
        baccDate.pack(side='bottom', padx = 5, pady=5)
        
    else:
        # убираются поля для ввода даты
        mb.showinfo("2", "2")
        
p1 = StringVar()
ch1 = Checkbutton(rootWork, text="Бронь",variable=p1, onvalue="Бронь есть", offvalue="Брони нет", command=isChecked)
p1.set("Брони нет")

def insertCB():
    cur.execute("Select ClientID, Surname, Name From Clients Order by ClientID")
    client=cur.fetchall()
    cb2['values']=client
    conn.commit()
    cur.execute("Select RoomID, Type From Rooms Where Employment = 'Свободно' Order By RoomID")
    room=cur.fetchall()
    cb1['values']=room

def select(event):
    global RoomID, ClientID, price, phone, bron
    s1=cb1.get()
    s=s1.find(" ")
    RoomID=cb1.get()[0:s]
    a1=cb2.get()
    a=a1.find(" ")
    ClientID = cb2.get()[0:a]
    cur.execute("Select Price_day From Rooms Where RoomID = ?;", [str(RoomID)])
    price = cur.fetchone()
    conn.commit()
    cur.execute("Select PhoneRoom From Rooms Where RoomID = ?;", [str(RoomID)])
    ph = cur.fetchone()
    if (ph[0]=="Есть"):
        cur.execute("Select Price_phone From Rooms Where RoomID = ?;", [str(RoomID)])
        phone=cur.fetchone()
    elif (ph[0]=="Нет"):
        phone=[0]
        
def addW(h1,h2,h3):
    try:
        pay = (price[0]*float(h3))+phone[0]
        Dnow = datetime.now()
        dateE = Dnow + timedelta(days=cd)
        answer = mb.askyesno(title="Подтверждение", message="К оплате "+str(pay)+" рублей, Подтвердить?")
        if answer:           
            if p1.get()=="Бронь есть":
                cur.execute("Insert into Work (RoomNum, ClientNum, DateStart, DateEnd, Duration_of_Days, Payment) values (?,?,?,?,?,?);",(h1,h2,ds.strftime("%d-%m-%Y"),de.strftime("%d-%m-%Y"),h3,pay))
                conn.commit()
                cur.execute("UPDATE Rooms Set Booking = 'Бронь есть', Employment = 'Занято' Where RoomID = ?;", [str(RoomID)])
                conn.commit()
                cur.execute("UPDATE Work Set Status = 'Не оплачено' Where WorkID = (Select WorkID From Work order by WorkID desc limit 1);")
                conn.commit()                              
                
            elif p1.get()=="Брони нет":
                cur.execute("Insert into Work (RoomNum, ClientNum, DateStart, DateEnd, Duration_of_Days, Payment) values (?,?,?,?,?,?);", (h1,h2,str(Dnow.strftime("%d-%m-%Y")),dateE.strftime("%d-%m-%Y"),h3,pay))
                conn.commit()
                cur.execute("UPDATE Rooms Set Employment = 'Занято' Where RoomID = ?;", [str(RoomID)])
                conn.commit()
                
                doc = docx.Document()
                p = doc.add_paragraph("Чек")
                p.alignment = 1
                font = p.runs[0].font
                font.size = Pt(30)
                p2 = doc.add_paragraph("Комната №")
                p2.add_run(" "+h1).bold = True
                p2.alignment = 1
                font1 = p2.runs[0].font
                font1.size = Pt(30)
                
                p3 = doc.add_paragraph("Клиент №")
                p3.add_run(" "+h2).bold = True
                p3.alignment = 1
                font2 = p3.runs[0].font
                font2.size = Pt(30)
                
                p4 = doc.add_paragraph("Количество дней - ")
                p4.add_run(" "+h3).bold = True
                p4.alignment = 1
                font3 = p4.runs[0].font
                font3.size = Pt(30)
                
                p5 = doc.add_paragraph("К оплате - ")
                p5.add_run(" "+str(pay)).bold = True
                p5.alignment = 2
                font4 = p5.runs[0].font
                font4.size = Pt(30)

                now = datetime.now()


                datenow = doc.add_paragraph(str(now.strftime("%d-%m-%Y")))
                font5 = datenow.runs[0].font

                datenow.alignment = 2
                font5.size = Pt(24)
                
                timenow = doc.add_paragraph(str(now.strftime("%H:%M")))
                timenow.alignment = 2
                font6 = timenow.runs[0].font
                font6.size = Pt(20)
                
                doc.save('Чек.docx')
                
    except:
        mb.showinfo("Ошибка!", "Заполните данные")
cb1.bind('<<ComboboxSelected>>', select)
cb2.bind('<<ComboboxSelected>>', select)

l1.pack(padx=5,pady=5)
cb1.pack(padx=5,pady=5)

l2.pack(padx=5,pady=5)
cb2.pack(padx=5,pady=5)

l3.pack(padx=5,pady=5)
e1.pack(padx=5,pady=5)

ch1.pack(padx=5,pady=5)

insertCB()
btnAc = Button(rootWork, text="Оформить", command=lambda:addW(h1=RoomID, h2=ClientID, h3=e1.get())).pack(padx=5,pady=5)

rootWork.mainloop()
