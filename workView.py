from tkinter import*
from tkinter import ttk
from tkinter import messagebox as mb
import sqlite3 as sql
import docx
from docx.shared import Pt
from datetime import datetime
from django.core.paginator import Paginator

conn=sql.connect('Hotel.db')
cur=conn.cursor()

rootView = Tk()
rootView.title("Просмотр аренды")

treev=ttk.Treeview(rootView)
treev.pack(side='left', pady=45)

scrlbar = ttk.Scrollbar(rootView, orient='vertical', command=treev.yview)
scrlbar.pack(side='right', fill='y')
treev.configure(yscrollcommand=scrlbar.set)

treev["columns"] = ("1", "2", "3", "4", "5", "6", "7", "8")

treev['show']='headings'

treev.column("1", width=110, anchor='c')
treev.column("2", width=120, anchor='c')
treev.column("3", width=120, anchor='c')
treev.column("4", width=120, anchor='se')
treev.column("5", width=100, anchor='se')
treev.column("6", width=120, anchor='se')
treev.column("7", width=80, anchor='se')
treev.column("8", width=100, anchor='se')

treev.heading("1", text="Номер комнаты")
treev.heading("2", text="ФИО клиента")
treev.heading("3", text="Дата заселения")
treev.heading("4", text="Дата выселения")
treev.heading("5", text="Количество дней")
treev.heading("6", text="Статус")
treev.heading("7", text="К оплате")
treev.heading("8", text="Статус заселения")




def view():
    global l, page, p
    treev.delete(*treev.get_children())
    cur.execute("Select RoomNum, Surname || ' ' || Name, DateStart, DateEnd, Duration_of_Days, Status, Payment, EmploymentStatus FROM Work inner join Clients ON ClientNum = ClientID order by RoomNum")
    obj = cur.fetchall()
    conn.commit()
    p = Paginator(obj, l)
    page = p.page(1)
    count = page.object_list
    for i in count:
        treev.insert("", 'end', values=i)
    conn.commit()
    
frame = LabelFrame(rootView, text='Поиск')
cbS = ttk.Combobox(frame, values=['Все','По Номеру комнаты', 'По ФИО клиента', 'По Количеству дней', 'К оплате'], state='readonly')
cbS.current(0)

def select(event):
    global stat, id_room, zasel, price, countD, FIClient
    it = treev.selection()[0]
    values=treev.item(it, option="values")
    id_room=values[0]
    FIClient = values[1]
    countD = values[4]
    price = values[6]
    stat = values[5]
    zasel = values[7]
    rootView.mainloop()

treev.bind('<<TreeviewSelect>>', select)

def search(*args):
    value=var.get()
    if value !="":
        if cbS.get()=="По ФИО клиента":
            cur.execute("select RoomNum, Surname || ' ' || Name, DateStart, DateEnd, Duration_of_Days, Status, Payment, EmploymentStatus FROM Work inner join Clients ON ClientNum = ClientID Where Surname || ' ' || Name Like ?",('{}%'.format(entrS.get()),))
            al = cur.fetchall()
            treev.delete(*treev.get_children())
            for i in al:
                treev.insert("", 'end', values=i)                
        if value.isdigit()==True:
            if cbS.get()!="По ФИО клиента":
                if cbS.get()=="По Номеру комнаты":
                    cur.execute("select RoomNum, Surname || ' ' || Name, DateStart, DateEnd, Duration_of_Days, Status, Payment, EmploymentStatus FROM Work inner join Clients ON ClientNum = ClientID Where RoomNum Like ?",('{}%'.format(entrS.get()),))
                    al = cur.fetchall()
                    treev.delete(*treev.get_children())
                    for i in al:
                        treev.insert("", 'end', values=i)
                
                elif cbS.get()=="По Количеству дней":
                    cur.execute("select RoomNum, Surname || ' ' || Name, DateStart, DateEnd, Duration_of_Days, Status, Payment, EmploymentStatus FROM Work inner join Clients ON ClientNum = ClientID Where Duration_of_Days Like ?",('{}%'.format(entrS.get()),))
                    al = cur.fetchall()
                    treev.delete(*treev.get_children())
                    for i in al:
                        treev.insert("", 'end', values=i)
                elif cbS.get()=="К оплате":
                    cur.execute("select RoomNum, Surname || ' ' || Name, DateStart, DateEnd, Duration_of_Days, Status, Payment, EmploymentStatus FROM Work inner join Clients ON ClientNum = ClientID Where Payment Like ?",('{}%'.format(entrS.get()),))
                    al = cur.fetchall()
                    treev.delete(*treev.get_children())
                    for i in al:
                        treev.insert("", 'end', values=i)
                elif cbS.get()=="Все":
                    view()
                    entrS.delete(0, 'end')
            else:
                mb.showinfo("Ошибка", "Здесь должно быть число")
                entrS.delete(0, 'end')
    else:
        view()
var=StringVar(rootView)
var.trace('w', search)

def retu():
    global stat, id_room
    try:
        if (stat=="Не оплачено"):
            answer = mb.askyesno(title="Подтверждение", message="Вы действительно хотите оплатить эту бронь?")
            if answer:
                cur.execute("UPDATE Work SET Status = 'Оплачено' Where RoomNum =?;", [str(id_room)])
                conn.commit()
                view()
        else:
            answer = mb.showinfo(title="Предупреждение", message="Эта бронь уже оплачена, выберите другую")
    except:
        mb.showinfo("Ошибка!", "Сначало выберите запись")


def left():
    global page, p
    if page.has_previous()==True:
        nump = page.previous_page_number()
        page = p.page(nump)
        treev.delete(*treev.get_children())
        count = page.object_list
        for i in count:
            treev.insert("", 'end', values=i)
    else:
        mb.showinfo("Внимание", "Больше назад нельзя")
count = Label(rootView)

def count_label():
    global countEnd
    cur.execute("select count(RoomNum) from Work")
    counEnd = cur.fetchone()
    conn.commit()
    count['text']="Всего "+str(counEnd[0]) +" записей"

def rigth():
    global page, p
    if page.has_next()==True:
        nump = page.next_page_number()
        page = p.page(nump)
        treev.delete(*treev.get_children())
        count = page.object_list
        for i in count:
            treev.insert("", 'end', values=i)
    else:
        mb.showinfo("Внимание", "Дальше нет страниц")
    
def CbPagInsert(event):
    global lim, l
    lim=cbL.get()
    l=cbL.get()
    view()

def createCheck(id_room, price, countD, FIClient):
    doc = docx.Document()
    p = doc.add_paragraph("Чек")
    p.alignment = 1
    font = p.runs[0].font
    font.size = Pt(30)
    p2 = doc.add_paragraph("Комната №")
    p2.add_run(" "+id_room).bold = True
    p2.alignment = 1
    font1 = p2.runs[0].font
    font01 = p2.runs[1].font
    font1.size = Pt(30)
    font01.size = Pt(30)
                
    p3 = doc.add_paragraph("Клиент №")
    p3.add_run(" "+FIClient).bold = True
    p3.alignment = 1
    font2 = p3.runs[0].font
    font02 = p3.runs[1].font
    font2.size = Pt(30)
    font02.size = Pt(30)
                
    p4 = doc.add_paragraph("Количество дней - ")
    p4.add_run(" "+countD).bold = True
    p4.alignment = 1
    font3 = p4.runs[0].font
    font03 = p4.runs[0].font
    font3.size = Pt(30)
    font03.size = Pt(30)
                
    p5 = doc.add_paragraph("К оплате - ")
    p5.add_run(" "+str(price)).bold = True
    p5.alignment = 2
    font4 = p5.runs[0].font
    font04 = p5.runs[1].font
    font4.size = Pt(30)
    font04.size = Pt(30)
    
    now = datetime.now()


    datenow = doc.add_paragraph(str(now.strftime("%d-%m-%Y")))
    font5 = datenow.runs[0].font

    datenow.alignment = 2
    font5.size = Pt(17)
                
    timenow = doc.add_paragraph(str(now.strftime("%H:%M")))
    timenow.alignment = 2
    font6 = timenow.runs[0].font
    font6.size = Pt(17)
                
    doc.save('Чек.docx')
    mb.showinfo("Успешно!", "Чек находится в корневой папке")

def zas():
    global id_room, zasel, price, countD, FIClient
    try:
        if zasel != "Заселён":
            answer = mb.askyesno(title="Подтверждение", message="Вы действительно хотите заселить этот номер?")
            if answer:
                cur.execute("Update Work Set EmploymentStatus = 'Заселён' Where RoomNum = ?", [str(id_room)])
                conn.commit()
                cur.execute("Update Rooms Set Employment = 'Занят' Where RoomID = ?", [str(id_room)])
                conn.commit()
                createCheck(id_room, price, countD, FIClient)
                view()
            
        else:
            mb.showinfo("Внимание", "Номер уже заселён, выберите другой")
    except:
        mb.showinfo("Ошибка", "Сначала необходимо выбрать запись")

cbL = ttk.Combobox(rootView, values=['5', '10','15'], state='readonly')
cbL.current(0)
cbL.bind('<<ComboboxSelected>>', CbPagInsert)
cbL.place(relx=0.2, rely=0.05)
treev.bind('<<TreeviewSelect>>', select)

lim=cbL.get()
l=cbL.get()

b_left=Button(rootView, text="<-", command=left).place(relx=0.15, rely=0.9)
b_right=Button(rootView, text="->", command=rigth).place(relx=0.30, rely=0.9)

count.place(relx=0.04, rely=0.03)

entrS = Entry(frame, textvariable=var)
ret = Button(frame, text='Оплатить', command=retu).pack(side='bottom', pady=5)
zas = Button(frame, text='Заселить', command=zas).pack(side='bottom', pady=5)
cbS.bind('<<ComboboxSelected>>', search)
cbS.pack(side='top', padx=5, pady=5)
entrS.pack(side='top', padx=5, pady=5)
frame.pack(side='top', padx=5, pady=5)
count_label()

view()


